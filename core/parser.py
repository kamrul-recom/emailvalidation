"""
Universal File Parser — extracts email addresses from every supported format.

Supported:
  - CSV (.csv)
  - Excel (.xlsx, .xls)
  - PDF — text-based AND scanned/image (OCR via tesseract)
  - Plain text (.txt)
  - Word document (.docx)
  - OpenDocument (.ods, .odt)
  - JSON (.json)
  - XML / HTML (.xml, .html, .htm)
  - ZIP archive (.zip) — recursively parses contained files
  - Direct paste (raw string via API)
"""

import re
import io
import csv
import json
import zipfile
from pathlib import Path

import openpyxl
import pdfplumber
import docx as python_docx
from lxml import etree

EMAIL_RE = re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}')


def _find(text: str) -> list[str]:
    return EMAIL_RE.findall(str(text or ""))


# ── Format parsers ────────────────────────────────────────────────────────────

def parse_csv(data: bytes) -> list[str]:
    text = data.decode("utf-8", errors="replace")
    emails = []
    try:
        # Try comma, then semicolon, then tab
        for delim in (",", ";", "\t", "|"):
            try:
                reader = csv.reader(io.StringIO(text), delimiter=delim)
                rows = list(reader)
                if len(rows) > 1 and len(rows[0]) > 0:
                    for row in rows:
                        for cell in row:
                            emails.extend(_find(cell))
                    if emails:
                        return emails
            except Exception:
                continue
    except Exception:
        pass
    # Fallback: raw regex scan
    return _find(text)


def parse_xlsx(data: bytes) -> list[str]:
    emails = []
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                emails.extend(_find(cell.value))
    wb.close()
    return emails


def parse_xls(data: bytes) -> list[str]:
    """Legacy .xls via openpyxl xlrd fallback or xlrd directly."""
    try:
        import xlrd
        wb = xlrd.open_workbook(file_contents=data)
        emails = []
        for sheet in wb.sheets():
            for r in range(sheet.nrows):
                for c in range(sheet.ncols):
                    emails.extend(_find(str(sheet.cell_value(r, c))))
        return emails
    except ImportError:
        # xlrd not installed — try openpyxl (works for some xls)
        try:
            return parse_xlsx(data)
        except Exception:
            return []


def parse_pdf(data: bytes) -> list[str]:
    """
    Two-stage PDF parser:
    1. pdfplumber for text-based PDFs (fast)
    2. pytesseract OCR for scanned/image PDFs (slower)
    """
    emails = []

    # Stage 1 — text extraction
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                emails.extend(_find(page.extract_text()))
                for table in (page.extract_tables() or []):
                    for row in (table or []):
                        for cell in (row or []):
                            emails.extend(_find(cell))
    except Exception:
        pass

    if emails:
        return emails

    # Stage 2 — OCR fallback for scanned PDFs
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(data, dpi=200)
        for img in images:
            text = pytesseract.image_to_string(img)
            emails.extend(_find(text))
    except Exception:
        pass

    return emails


def parse_txt(data: bytes) -> list[str]:
    text = data.decode("utf-8", errors="replace")
    return _find(text)


def parse_docx(data: bytes) -> list[str]:
    emails = []
    doc = python_docx.Document(io.BytesIO(data))
    # Paragraphs
    for para in doc.paragraphs:
        emails.extend(_find(para.text))
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                emails.extend(_find(cell.text))
    # Headers and footers
    for section in doc.sections:
        for hdr in [section.header, section.footer]:
            if hdr:
                for para in hdr.paragraphs:
                    emails.extend(_find(para.text))
    return emails


def parse_ods(data: bytes) -> list[str]:
    """OpenDocument Spreadsheet (.ods)"""
    emails = []
    try:
        from odf.opendocument import load as odf_load
        from odf.table import Table, TableRow, TableCell
        from odf.text import P
        doc = odf_load(io.BytesIO(data))
        for sheet in doc.spreadsheet.getElementsByType(Table):
            for row in sheet.getElementsByType(TableRow):
                for cell in row.getElementsByType(TableCell):
                    for p in cell.getElementsByType(P):
                        emails.extend(_find(str(p)))
    except Exception:
        pass
    return emails


def parse_odt(data: bytes) -> list[str]:
    """OpenDocument Text (.odt)"""
    emails = []
    try:
        from odf.opendocument import load as odf_load
        from odf.text import P
        doc = odf_load(io.BytesIO(data))
        for p in doc.text.getElementsByType(P):
            emails.extend(_find(str(p)))
    except Exception:
        pass
    return emails


def parse_json(data: bytes) -> list[str]:
    """JSON — recursively walks all string values."""
    def walk(obj):
        found = []
        if isinstance(obj, str):
            found.extend(_find(obj))
        elif isinstance(obj, dict):
            for v in obj.values():
                found.extend(walk(v))
        elif isinstance(obj, (list, tuple)):
            for item in obj:
                found.extend(walk(item))
        return found

    text = data.decode("utf-8", errors="replace")
    try:
        obj = json.loads(text)
        return walk(obj)
    except json.JSONDecodeError:
        # Partial / JSONL — scan line by line
        emails = []
        for line in text.splitlines():
            try:
                emails.extend(walk(json.loads(line)))
            except Exception:
                emails.extend(_find(line))
        return emails


def parse_xml_html(data: bytes) -> list[str]:
    """XML and HTML — strips tags, extracts text content + attribute values."""
    emails = []
    text_raw = data.decode("utf-8", errors="replace")
    # Raw regex first (catches emails in attributes like mailto:)
    emails.extend(_find(text_raw))
    # Then lxml text extraction for clean content
    try:
        root = etree.fromstring(data)
        for node in root.iter():
            emails.extend(_find(node.text))
            emails.extend(_find(node.tail))
            for val in node.attrib.values():
                emails.extend(_find(val))
    except Exception:
        try:
            from lxml import html as lhtml
            root = lhtml.fromstring(data)
            emails.extend(_find(root.text_content()))
        except Exception:
            pass
    return emails


def parse_zip(data: bytes) -> list[str]:
    """ZIP — recursively parse every file inside."""
    emails = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for name in zf.namelist():
                if name.endswith("/"):
                    continue
                try:
                    inner_bytes = zf.read(name)
                    inner_emails, _ = parse_file(name, inner_bytes)
                    emails.extend(inner_emails)
                except Exception:
                    pass
    except Exception:
        pass
    return emails


def parse_raw_text(text: str) -> list[str]:
    """Direct paste / API — accepts a raw string."""
    return _find(text)


# ── Router ────────────────────────────────────────────────────────────────────

PARSERS = {
    ".csv":  parse_csv,
    ".xlsx": parse_xlsx,
    ".xls":  parse_xls,
    ".pdf":  parse_pdf,
    ".txt":  parse_txt,
    ".docx": parse_docx,
    ".doc":  parse_docx,
    ".ods":  parse_ods,
    ".odt":  parse_odt,
    ".json": parse_json,
    ".jsonl":parse_json,
    ".xml":  parse_xml_html,
    ".html": parse_xml_html,
    ".htm":  parse_xml_html,
    ".zip":  parse_zip,
}

SUPPORTED_EXTENSIONS = sorted(PARSERS.keys())


def parse_file(filename: str, data: bytes) -> tuple[list[str], str]:
    """
    Returns (emails, error_message).
    error_message is "" on success.
    """
    ext = Path(filename).suffix.lower()
    parser = PARSERS.get(ext)
    if not parser:
        # Unknown extension — try plain text fallback
        return parse_txt(data), f"Unknown format '{ext}', scanned as plain text"
    try:
        emails = parser(data)
        return emails, ""
    except Exception as e:
        return [], f"Failed to parse {filename}: {e}"
